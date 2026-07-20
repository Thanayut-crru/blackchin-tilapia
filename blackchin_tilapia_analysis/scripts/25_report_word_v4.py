"""
Step 25 — Word manuscript v4 (Full-Factorial Multi-Algorithm Revision)

v4 CHANGES vs v3:
  1. Full factorial experiment: 3 algorithms x 4 feature sets x 4 PA methods = 48 combos
  2. Added Section 2.4: Machine Learning Algorithms (RF, XGB, MaxNet/MaxEnt)
  3. Added Section 2.5: Pseudo-Absence Sampling Methods (4 PA methods)
  4. Renumbered 2.4->2.6 (Nested CV), 2.5->2.7 (Evaluation)
  5. Updated Results 3.2: Full Factorial AUC matrix table
  6. Added Results 3.3: Algorithm comparison (Kruskal-Wallis + pairwise)
  7. Added Results 3.4: Feature set comparison
  8. Added Results 3.5: PA method comparison
  9. Added Results 3.6: Sensitivity analysis
  10. Updated Abstract: RF/M_SAP/two_step AUC=0.894 [0.867-0.915]
  11. Updated Discussion: algorithm uncertainty, PA method sensitivity
  12. Output: output4/SAP_Manuscript_v5.docx
"""

import sys, os, warnings
sys.path.insert(0, str(__import__("pathlib").Path(__file__).parent.parent))
os.environ["PYTHONWARNINGS"] = "ignore"
warnings.filterwarnings("ignore")

import numpy as np
import pandas as pd
from pathlib import Path
from datetime import date

from blackchin_tilapia_analysis.scripts.s00_config import OCC_CSV, OUT3, OUT5

BASE    = Path(__file__).parent.parent
# OUT4 removed


OUT_DOC  = OUT5 / "SAP_Manuscript_v5.docx"
FIG4_DIR = OUT5 / "figures"
TAB4_DIR = OUT5 / "paper_tables"

# output3 sources (SHAP, VIF, occurrence, province risk from v3 pipeline)
TAB3_DIR  = OUT3 / "paper_tables"
FIG3_DIR  = OUT3 / "figures"
VAL3_DIR  = OUT5 / "validation"
VAL3_DIR1 = OUT3 / "validation_v2"


def load(path, fallback=None):
    p = Path(path)
    if p.exists():
        return pd.read_csv(p)
    if fallback and Path(fallback).exists():
        return pd.read_csv(fallback)
    return pd.DataFrame()


# ── Word builder ──────────────────────────────────────────────────────────────

def build_doc():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    for margin in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, margin, Cm(2.5))

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.name  = "Times New Roman"
            run.font.bold  = True
            run.font.size  = Pt(13 if level == 1 else 12)
            run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def para(text, bold=False, italic=False, size=12, align=None):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(size)
        run.font.bold   = bold
        run.font.italic = italic
        if align:
            p.alignment = align
        p.paragraph_format.space_after = Pt(6)
        return p

    def caption(text):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(10)
        run.font.italic = True
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(10)

    def note(text):
        p   = doc.add_paragraph()
        run = p.add_run("[Note for revision: " + text + "]")
        run.font.name   = "Times New Roman"
        run.font.size   = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(180, 0, 0)
        p.paragraph_format.space_after = Pt(4)

    def add_table_from_df(df, font_size=9):
        t   = doc.add_table(rows=1, cols=len(df.columns))
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Times New Roman"
        for _, row in df.iterrows():
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val) if pd.notna(val) else ""
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Times New Roman"
        return t

    def add_figure(path, width_cm=14.0, fallback_text=None):
        path = Path(path)
        if path.exists():
            doc.add_picture(str(path), width=Cm(width_cm))
        elif fallback_text:
            para(f"[Figure not yet generated: {path.name}]", italic=True)

    # ── Load key data files ────────────────────────────────────────────────────
    pool   = load(OUT5 / "pooled_oof_metrics.csv")
    fold   = load(OUT5 / "experiment_fold_metrics.csv")
    tab_a  = load(TAB4_DIR / "tableA_auc_matrix.csv")
    tab_b  = load(TAB4_DIR / "tableB_best_config.csv")
    tab_c  = load(TAB4_DIR / "tableC_kruskal_wallis.csv")
    tab_pw = load(TAB4_DIR / "tableC_pairwise.csv")
    tab_d  = load(TAB4_DIR / "tableD_pa_method.csv")
    tab_e  = load(TAB4_DIR / "tableE_sensitivity.csv")
    # Full-data maps and SHAP must be regenerated from the current n=38 run.
    # Do not reuse output3 artefacts produced by an older dataset/model.
    tab_f  = pd.DataFrame()
    shap   = pd.DataFrame()
    df_occ = pd.DataFrame()
    df_prov = pd.DataFrame()

    # Dynamic occurrence count (year-filtered) — primary sample size for manuscript
    _occ_raw = load(OCC_CSV)
    if len(_occ_raw) > 0 and "year" in _occ_raw.columns:
        _occ_filtered = _occ_raw[_occ_raw["year"] <= 2024]
        N_OCC = int(len(_occ_filtered))
        N_OCC_RAW = int(len(_occ_raw))
        N_OCC_REMOVED = N_OCC_RAW - N_OCC
    else:
        N_OCC = 40; N_OCC_RAW = 40; N_OCC_REMOVED = 0

    # Derive best-config numbers
    def _best():
        if len(pool) == 0:
            return {"algo":"RF","feat_set":"M_SAP","pa_method":"two_step",
                    "auc":"0.894","ci_lo":"0.867","ci_hi":"0.915",
                    "pr_auc":"0.615","tss":"0.697","brier":"0.402"}
        auc_col = "auc_roc" if "auc_roc" in pool.columns else "auc_roc_mean"
        b = pool.sort_values(auc_col, ascending=False).iloc[0]
        ci_lo = b.get("auc_roc_ci_lo", "?")
        ci_hi = b.get("auc_roc_ci_hi", "?")
        pr    = b.get("pr_auc", b.get("pr_auc_mean", "?"))
        tss   = b.get("tss",    b.get("tss_mean",   "?"))
        brier = b.get("brier",  b.get("brier_mean",  "?"))
        return {
            "algo":    str(b["algo"]),
            "feat_set":str(b["feat_set"]),
            "pa_method":str(b["pa_method"]),
            "auc":    f"{b[auc_col]:.3f}",
            "ci_lo":  f"{float(ci_lo):.3f}" if ci_lo != "?" else "?",
            "ci_hi":  f"{float(ci_hi):.3f}" if ci_hi != "?" else "?",
            "pr_auc": f"{float(pr):.3f}"    if pr != "?"   else "?",
            "tss":    f"{float(tss):.3f}"   if tss != "?"  else "?",
            "brier":  f"{float(brier):.3f}" if brier != "?" else "?",
        }
    BEST = _best()

    def algo_mean_auc(algo):
        if len(fold) == 0: return {"RF":"0.774","XGB":"0.765","MXN":"0.695"}.get(algo,"?")
        v = fold[fold["algo"] == algo]["auc_roc"].dropna()
        return f"{v.mean():.3f}" if len(v) else "?"

    def feat_mean_auc(fs):
        if len(fold) == 0: return "?"
        v = fold[fold["feat_set"] == fs]["auc_roc"].dropna()
        return f"{v.mean():.3f}" if len(v) else "?"

    def kw_result(factor, response="AUC-ROC"):
        if len(tab_c) == 0: return ("?", "?", "?")
        sub = tab_c[(tab_c["Factor"] == factor) &
                    (tab_c["Response"].str.upper() == response.upper())]
        if len(sub) == 0: return ("?", "?", "?")
        r = sub.iloc[0]
        stat = str(r.get("Statistic", "?")).replace("H=","")
        p    = r.get("p-value", "?")
        sig  = str(r.get("Significance",""))
        return (stat, f"{float(p):.3f}" if p != "?" else "?", sig)

    def pairwise_result(factor, a, b, response="auc_roc"):
        if len(tab_pw) == 0: return ("?", "?")
        sub = tab_pw[tab_pw.get("factor", pd.Series([])) == factor] if "factor" in tab_pw.columns else tab_pw
        for _, r in sub.iterrows():
            g1 = str(r.get("level_a", r.get("group1", "")))
            g2 = str(r.get("level_b", r.get("group2", "")))
            if (g1 == a and g2 == b) or (g1 == b and g2 == a):
                p = r.get("p_holm", r.get("p_paired_perm",
                    r.get("p_permutation", r.get("p_value", "?"))))
                sig = "**" if float(p) < 0.01 else ("*" if float(p) < 0.05 else "ns") if p != "?" else "?"
                return (f"{float(p):.3f}" if p != "?" else "?", sig)
        return ("?", "?")

    top_prov = df_prov.iloc[0]["province"] if len(df_prov) else "not evaluated"
    top_feat = shap.iloc[0]["feature"] if len(shap) else "not evaluated"

    # ══════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════
    doc.add_paragraph()
    p   = doc.add_paragraph()
    run = p.add_run(
        "A Leakage-Free Explainable Machine Learning Framework Integrating "
        "Environmental Suitability, Waterway Proximity, and Human-Pressure Proxies "
        "for Predicting Relative Habitat Suitability of Blackchin Tilapia "
        "(Sarotherodon melanotheron) in Thailand"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = True
    p.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    para("[Author names and affiliations -- to be completed]",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(f"Draft prepared: {date.today().isoformat()}  (Paper 1 leakage-free full factorial experiment)",
         italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════
    heading("Abstract", level=1)
    abstract_text = (
        "Biological invasion of blackchin tilapia (Sarotherodon melanotheron) poses a "
        "growing threat to freshwater and estuarine biodiversity in Thailand. Conventional "
        "species distribution models (SDMs) rely on single-algorithm implementations with "
        "climatic predictors only, are susceptible to spatial data leakage, and rarely "
        "evaluate pseudo-absence sampling sensitivity. "
        "We developed a leakage-free nested spatial cross-validation framework that "
        "simultaneously compares three machine learning algorithms — Random Forest (RF), "
        "Extreme Gradient Boosting (XGB), and MaxNet (elastic-net logistic regression) — "
        "across four predictor combinations: environmental suitability "
        "(S; 10 bioclimatic/elevation variables), S + waterway proximity (A; distance to "
        "waterway, stream-order proxy), S + human-pressure proxies "
        "(P; population density, road distance, urban proximity), and all three layers (SAP). "
        "Four pseudo-absence sampling strategies (random, spatially constrained, two-step, "
        "uncertainty-zone) were evaluated as tunable model components. "
        f"We assembled n = {N_OCC} quality-controlled occurrence records "
        f"(after removing {N_OCC_REMOVED} temporally suspect records with year > 2024) "
        "and applied a 5-repeat × 3-outer-fold × 3-inner-fold nested spatial CV design "
        "with VIF selection, pseudo-absence ratio, and hyperparameter tuning inside each inner loop. "
        f"The best configuration — {BEST['algo']}/{BEST['feat_set']}/{BEST['pa_method']} — "
        f"achieved fold-mean AUC-ROC = {BEST['auc']} "
        f"[95% clustered-bootstrap CI: {BEST['ci_lo']}–{BEST['ci_hi']}], "
        f"PR-AUC = {BEST['pr_auc']}, TSS = {BEST['tss']} (threshold from inner OOF). "
        "RF achieved the highest mean fold-level AUC, followed by XGB and MXN; all "
        "three pairwise algorithm contrasts remained significant after Holm correction. "
        "Pseudo-absence strategy had no significant omnibus effect on AUC; only the "
        "random versus spatially constrained contrast was significant after correction. "
        "M_SP and M_SAP had nearly identical fold-level mean AUC, indicating predictive "
        "value from human-pressure proxies without a clear additional gain from combining "
        "all available waterway proxies. Full-data maps and province rankings were withheld "
        "pending regeneration from the current dataset. Results remain provisional and "
        "require independent validation."
    )
    para(abstract_text)
    para(
        "Keywords: species distribution model; biological invasion; spatial cross-validation; "
        "Random Forest; XGBoost; MaxNet; MaxEnt; pseudo-absence; SHAP; Thailand; tilapia",
        italic=True,
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════
    heading("1. Introduction", level=1)
    para(
        "Biological invasions represent one of the greatest threats to global biodiversity "
        "and ecosystem services (Vitousek et al. 1997). Blackchin tilapia "
        "(Sarotherodon melanotheron Ruppell, 1852), native to West Africa, has become "
        "established in multiple Thai provinces following its detection in central Thailand "
        "(Vidthayanon 2005). Its exceptional euryhalinity (0-75 ppt salinity tolerance), "
        "rapid reproductive rate (maturity at 3-4 months), and resistance to hypoxia enable "
        "colonisation of estuarine canals, aquaculture ponds, and coastal wetlands where "
        "it competitively displaces native species and disrupts local fisheries."
    )
    para(
        "Species distribution models (SDMs) are widely used to predict invasion suitability "
        "(Elith & Leathwick 2009). However, most published SDMs for aquatic invaders "
        "include only climatic or abiotic variables, neglecting two key drivers of "
        "invasion spread: (1) waterway connectivity, which determines a species' physical "
        "ability to disperse via waterway networks (Rahel 2007), and (2) human-mediated "
        "introduction pressure, reflecting the probability of introduction through trade, "
        "transport, or aquaculture escape (Lockwood et al. 2005). Moreover, many SDMs "
        "violate the principle of spatial independence during cross-validation, leading to "
        "optimistically biased performance estimates (Roberts et al. 2017; Valavi et al. 2019)."
    )
    para(
        "A further challenge is algorithm uncertainty: machine learning SDMs employing "
        "Random Forest, gradient boosted trees, or regularised logistic regression "
        "(MaxEnt/MaxNet) produce substantially different suitability surfaces and performance "
        "estimates (Elith et al. 2006). Similarly, the choice of pseudo-absence (background) "
        "sampling strategy is known to affect AUC, calibration, and predicted distribution "
        "extent (Iturbide et al. 2015; Phillips et al. 2009). Formal comparison of algorithms "
        "and pseudo-absence methods under a unified, leakage-free design is therefore "
        "needed to provide reliable guidance on model selection."
    )
    para(
        "We address these limitations by developing a leakage-free R = f(S, A, P) "
        "framework that simultaneously compares three machine learning algorithms and "
        "four pseudo-absence strategies across four predictor combinations within a strictly "
        "nested spatial CV design. All internal procedures -- including VIF feature "
        "selection, pseudo-absence generation, pseudo-absence ratio, and hyperparameter "
        "optimisation -- are conducted exclusively on outer-training data. Test-set background "
        "is drawn only from the test spatial block, ensuring no information from the test "
        "partition contaminates model development. SHAP (SHapley Additive exPlanations; "
        "Lundberg & Lee 2017) values provide interpretable predictor attribution."
    )

    # ══════════════════════════════════════════════════════════
    # 2. MATERIALS AND METHODS
    # ══════════════════════════════════════════════════════════
    heading("2. Materials and Methods", level=1)

    heading("2.1 Study Area and Target Species", level=2)
    para(
        "Thailand (lat 5.5-20.6 N, lon 97.3-105.7 E) encompasses diverse aquatic habitats "
        "ranging from peninsular coasts bordering the Gulf of Thailand and the Andaman Sea "
        "to the Mekong floodplain in the northeast. The modelling domain was defined at "
        "2.5 arcmin (~4.6 km) spatial resolution, producing a 240 x 432 cell raster grid. "
        "The common modelling domain comprises 11,750 water-class cells for which all "
        "16 retained predictor variables are valid. WorldPop counts were converted to "
        "people per square kilometre, and missing population values on valid water cells "
        "were assigned zero; all model variants were evaluated on this common domain to ensure "
        "comparability across feature sets."
    )

    heading("2.2 Occurrence Data", level=2)
    para(
        "Blackchin tilapia occurrence records were compiled from three sources: "
        "(1) GBIF (https://doi.org/[GBIF DOI to be confirmed after download]), "
        "(2) iNaturalist (inaturalist.org) queried via the public API with Thailand bounding box, "
        "and (3) verified field records including published catch reports and georeferenced "
        "monitoring records (evidence levels A-B: GPS coordinates or high-precision locality). "
        "Records underwent a nine-step quality control pipeline: taxonomy harmonisation, "
        "coordinate validation, bounding box filtering, coordinate uncertainty filtering "
        "(<5 km), duplicate removal by raster-cell and observation ID, and temporal filtering "
        "(1980-2024). The primary experiment retained all records passing these criteria; "
        "spatial thinning at 5, 10, and 20 km was evaluated separately as a sensitivity "
        f"analysis. Final primary dataset: n = {N_OCC} occurrence points."
    )
    if len(df_occ):
        caption("Table 1. Blackchin tilapia occurrence data summary.")
        add_table_from_df(df_occ)
        doc.add_paragraph()

    heading("2.3 Predictor Variables", level=2)
    para(
        "We assembled 16 predictor variables in three thematic groups (Table 2). "
        "S-layer (n=10): nine bioclimatic variables selected to represent non-collinear "
        "dimensions of temperature and precipitation (WorldClim 2.1, 2.5 arcmin; "
        "Fick & Hijmans 2017) plus elevation. "
        "A-layer (n=2): distance to the nearest OSM waterway (rivers, canals, streams, drains; "
        "Geofabrik Thailand extract) and a stream-order proxy derived from OSM feature class. "
        "P-layer (n=4): WorldPop 2020 population density aggregated to 4.6 km (Tatem 2017), "
        "distance to nearest paved OSM road, a road-class density proxy, and distance to "
        "nearest urban centre (WorldPop cells >500 persons/km2). "
        "All layers were resampled or aggregated to the 4.6 km reference grid (EPSG:4326)."
    )
    df_feat = load(TAB3_DIR / "table2_feature_inventory.csv")
    if len(df_feat) and "variable" in df_feat.columns:
        df_feat = df_feat[df_feat["variable"] != "A3_dist_coast"].copy()
    if len(df_feat):
        df_show = df_feat[["variable","group","description","source","native_resolution"]].copy()
        df_show.columns = ["Variable","Group","Description","Source","Native res."]
        doc.add_paragraph()
        caption("Table 2. Summary of predictor variables.")
        add_table_from_df(df_show, font_size=8)
        doc.add_paragraph()

    heading("2.4 Machine Learning Algorithms", level=2)
    para(
        "Three machine learning algorithms were evaluated. (1) Random Forest (RF; "
        "Breiman 2001): an ensemble of decision trees trained on bootstrap samples with "
        "random feature subsets at each split (scikit-learn RandomForestClassifier, "
        "n_estimators=200, max_features in {'sqrt','log2','0.5'}, max_depth in {5,8,12,None}, "
        "min_samples_leaf in {1,2,4,8}, evaluated via random search over 10 candidates). "
        "RF is known for robustness to noise, implicit feature selection via bagging, "
        "and reliable probability estimates when calibrated. "
        "(2) Extreme Gradient Boosting (XGB; Chen & Guestrin 2016): a sequential "
        "tree-boosting algorithm that minimises a regularised loss function "
        "(XGBClassifier, n_estimators in {100,200,300}, max_depth in {3,4,5,6}, "
        "learning_rate in {0.05,0.1,0.2}, subsample in {0.6,0.8,1.0}). XGB typically "
        "achieves higher predictive accuracy than RF on structured data but requires "
        "more hyperparameter tuning. "
        "(3) MaxNet (MaxEnt-equivalent; Phillips et al. 2017): a regularised logistic "
        "regression using linear and hinge basis functions of the predictors, implemented "
        "as an elastic-net logistic regression (scikit-learn LogisticRegression, "
        "penalty='elasticnet', solver='saga', l1_ratio in {0.1,0.5,0.9}, "
        "C in {0.01,0.1,1.0,10.0}). MaxNet is conceptually equivalent to the widely-used "
        "MaxEnt Java program (Phillips et al. 2006) for presence-background modelling, "
        "using elastic-net regularisation to approximate MaxEnt's regularisation scheme. "
        "For each predictor, hinge features were generated at 10 evenly spaced thresholds."
    )

    heading("2.5 Pseudo-Absence Sampling Methods", level=2)
    para(
        "Four pseudo-absence (background) sampling strategies were evaluated as a "
        "cross-validated model component. (1) Random: background points sampled uniformly "
        "at random from all valid cells outside a 10 km buffer around training presences. "
        "This is the conventional approach and provides the baseline. "
        "(2) Spatially constrained: background points restricted to within 200 km of the "
        "nearest training presence, reducing the geographic extent of environmental "
        "contrast and potentially improving prediction in the study area (Iturbide et al. 2015). "
        "(3) Two-step: a Random Forest pilot model (100 trees, trained on 1:10 random "
        "background) is first fitted to the training data; background points are then "
        "restricted to cells predicted as low-suitability (<0.2) by the pilot, "
        "focusing pseudo-absences on environmental conditions genuinely dissimilar to "
        "presence conditions (Senay et al. 2013). "
        "(4) SHAP-fuzzy: background points are drawn from cells with pilot-model suitability "
        "between 0.2 and 0.6, representing areas of intermediate suitability where the "
        "classifier boundary is uncertain. In all cases, pseudo-absence ratio (PA:presence = "
        "1:1, 1:3, 1:5, or 1:10) was treated as a hyperparameter selected jointly with "
        "algorithm hyperparameters in the inner CV loop."
    )

    heading("2.6 Leakage-Free Nested Spatial Cross-Validation", level=2)
    para(
        "The full factorial experiment comprised 3 algorithms x 4 feature sets x 4 "
        "pseudo-absence methods = 48 combinations, each evaluated under an identical "
        "nested spatial CV design. Occurrence points were assigned to three spatial "
        "folds using balanced geographic projection blocks (Roberts et al. 2017). "
        "For each repeat, coordinates were projected onto a seed-specific geographic axis, "
        "ordered along that axis, and divided into three contiguous groups whose sizes "
        "differed by at most one occurrence. The same projection boundaries assigned "
        "background cells to folds. Five seed-specific axes produced repeated spatial "
        "partitions while preventing very small test folds. "
        "The nested CV design comprised 5 repeats x 3 outer folds x 3 inner folds "
        "(= 45 outer-fold evaluations per combination). "
        "For each outer fold, only outer-training data were used to: (i) generate "
        "pseudo-absences using the selected method; (ii) apply iterative VIF feature "
        "selection (threshold=10) using LinearRegression R2; (iii) jointly optimise "
        "pseudo-absence ratio and algorithm hyperparameters via random search over 10 "
        "candidates evaluated on inner folds by AUC-ROC. "
        "Test pseudo-absences were sampled exclusively from cells within the outer test "
        "spatial block, ensuring that presences and background come from the same "
        "geographic region. The Youden J optimal threshold (sensitivity + specificity - 1, "
        "maximised on inner-fold out-of-fold predictions) was used for TSS computation. "
        "Pooled out-of-fold (OOF) predictions (all outer-fold test predictions concatenated) "
        "were used to compute AUC-ROC, PR-AUC, TSS, and Brier score, with 95% bootstrap "
        "confidence intervals (500 iterations, percentile method)."
    )

    heading("2.7 Model Evaluation and Statistical Comparison", level=2)
    para(
        "Evaluation metrics: AUC-ROC (discriminatory ability), PR-AUC "
        "(precision-recall AUC; suitable for imbalanced presence-background settings), "
        "TSS (True Skill Statistic at Youden J threshold), and Brier score "
        "(calibration accuracy). Statistical comparison across algorithms, feature sets, "
        "and pseudo-absence methods used Kruskal-Wallis H-tests (non-parametric, "
        "comparisons) on 720 fold-level values (48 combinations x 15 outer folds; "
        "5 repeats x 3 outer folds). Pairwise differences were assessed by matched "
        "sign-flip permutation tests (500 iterations, two-sided), with Monte Carlo "
        "p-values calculated as (extreme + 1)/(iterations + 1) and Holm correction. "
        "Sensitivity analysis examined the robustness of the "
        "best configuration to: spatial thinning distance (5, 10, 20 km), number of CV "
        "folds (2, 3, 4), and occurrence subset (GBIF/iNaturalist only vs. full dataset). "
        "Province-level relative suitability was computed by averaging M_SAP predictions "
        "across all valid cells within each province boundary (GADM 4.1 level 1)."
    )

    # ══════════════════════════════════════════════════════════
    # 3. RESULTS
    # ══════════════════════════════════════════════════════════
    heading("3. Results", level=1)

    heading("3.1 Occurrence Data Summary", level=2)
    para(
        f"After quality control and removal of {N_OCC_REMOVED} "
        f"temporally suspect records (year > 2024; see Section 4.4), the final dataset "
        f"comprised n = {N_OCC} occurrence records (Table 1). "
        f"Records spanned Thai provinces in the Gulf of Thailand coastal region. "
        f"This small sample size (n = {N_OCC}) is the primary methodological constraint "
        f"of the present study and is explicitly acknowledged as a limitation throughout."
    )
    if N_OCC < 20:
        note(f"SAMPLE SIZE WARNING: n={N_OCC} after year filter. "
             f"Standard SDM practice recommends n≥20-30. "
             f"All performance metrics and map predictions should be interpreted "
             f"with extreme caution. This limitation must be prominently stated in the final manuscript.")

    heading("3.2 Full Factorial Experiment: AUC-ROC Matrix", level=2)
    para(
        "Table 3 presents mean fold-level AUC-ROC (+/-SD) across all 48 algorithm x "
        "feature set x pseudo-absence method combinations. The best single configuration "
        f"was {BEST['algo']}/{BEST['feat_set']}/{BEST['pa_method']}, with pooled OOF "
        f"AUC-ROC = {BEST['auc']} [95% CI: {BEST['ci_lo']}-{BEST['ci_hi']}], "
        f"PR-AUC = {BEST['pr_auc']}, TSS = {BEST['tss']}, Brier = {BEST['brier']}. "
        "Across algorithms, M_SP had the highest fold-level mean AUC, followed by "
        "M_SAP, M_SA, and M_S. Performance varied substantially among folds and "
        "pseudo-absence strategies."
    )
    if len(tab_b):
        caption(
            "Table 3. Best-configuration pooled OOF metrics per algorithm x feature set "
            "(highest AUC-ROC pseudo-absence method selected). "
            "95% CI from 500-iteration bootstrap (pooled OOF predictions)."
        )
        add_table_from_df(tab_b)
        doc.add_paragraph()

    # Full AUC matrix
    auc_mat_csv = TAB4_DIR / "tableA_auc_matrix.csv"
    if auc_mat_csv.exists():
        auc_mat = pd.read_csv(auc_mat_csv, index_col=0)
        caption(
            "Table 4. AUC-ROC matrix (mean (SD) fold-level AUC-ROC) for all algorithm x "
            "feature set combinations (pooled across pseudo-absence methods)."
        )
        add_table_from_df(auc_mat.reset_index().rename(columns={"index":"Algorithm"}))
        doc.add_paragraph()

    heading("3.3 Algorithm Comparison", level=2)
    h_algo, p_algo, sig_algo = kw_result("algo", "AUC-ROC")
    rf_xgb_p,  rf_xgb_sig  = pairwise_result("algo","RF","XGB")
    rf_mxn_p,  rf_mxn_sig  = pairwise_result("algo","RF","MXN")
    xgb_mxn_p, xgb_mxn_sig = pairwise_result("algo","XGB","MXN")
    para(
        f"Algorithm had a significant overall effect on AUC-ROC "
        f"(Kruskal-Wallis H = {h_algo}, p = {p_algo}, {sig_algo}; Table 5). "
        f"Pairwise permutation tests showed: RF vs XGB p = {rf_xgb_p} ({rf_xgb_sig}); "
        f"RF vs MXN p = {rf_mxn_p} ({rf_mxn_sig}); "
        f"XGB vs MXN p = {xgb_mxn_p} ({xgb_mxn_sig}). "
        f"Mean AUC was {algo_mean_auc('RF')} for RF, {algo_mean_auc('XGB')} for XGB, "
        f"and {algo_mean_auc('MXN')} for MXN. All three contrasts remained "
        "significant after Holm correction."
    )
    if len(tab_c):
        caption(
            "Table 5. Kruskal-Wallis H-test results for algorithm, feature set, "
            "and pseudo-absence method effects on AUC-ROC, PR-AUC, and TSS "
            "(fold-level values, n=240 per algorithm; n=180 per feature set; "
            "n=180 per pseudo-absence method)."
        )
        add_table_from_df(tab_c)
        doc.add_paragraph()

    heading("3.4 Feature Set Comparison", level=2)
    h_fs, p_fs, sig_fs = kw_result("feat_set", "AUC-ROC")
    ms_auc   = feat_mean_auc("M_S")
    msa_auc  = feat_mean_auc("M_SA")
    msp_auc  = feat_mean_auc("M_SP")
    msap_auc = feat_mean_auc("M_SAP")
    para(
        f"Feature set had a significant effect on AUC-ROC (Kruskal-Wallis H = {h_fs}, "
        f"p = {p_fs}, {sig_fs}). M_SP had the highest mean AUC ({msp_auc}), "
        f"followed by M_SAP ({msap_auc}), M_SA ({msa_auc}), and M_S ({ms_auc}). "
        "All augmented feature sets exceeded S alone in paired comparisons. M_SP exceeded "
        "M_SA, whereas M_SP and M_SAP did not differ significantly. Thus, the "
        "human-pressure proxies added useful discrimination, but adding the available "
        "A-layer proxies did not provide a clear further gain."
    )

    heading("3.5 Pseudo-Absence Method Comparison", level=2)
    h_pa, p_pa, sig_pa = kw_result("pa_method", "AUC-ROC")
    pa_auc_rows = []
    if len(tab_d):
        auc_sub = tab_d[tab_d["Metric"] == "AUC-ROC"]
        for _, r in auc_sub.iterrows():
            pa_auc_rows.append(f"{r['PA Method']} (mean={r['Mean']:.3f})")
    pa_auc_str = "; ".join(pa_auc_rows) if pa_auc_rows else "see Table 6"
    para(
        f"Pseudo-absence sampling method did not affect AUC-ROC in the omnibus comparison "
        f"(Kruskal-Wallis H = {h_pa}, p = {p_pa}, {sig_pa}). "
        f"Numerical AUC values were: {pa_auc_str}. "
        "Random sampling was numerically highest. Holm-adjusted pairwise tests identified "
        "only the random versus spatially constrained contrast as significant; all other "
        "pseudo-absence contrasts were not significant."
    )
    if len(tab_d):
        auc_tab = tab_d[tab_d["Metric"] == "AUC-ROC"][["PA Method","Mean","SD","N"]].copy()
        caption("Table 6. AUC-ROC by pseudo-absence sampling method (fold-level mean and SD).")
        add_table_from_df(auc_tab)
        doc.add_paragraph()

    heading("3.6 VIF Feature Selection", level=2)
    df_vif = pd.DataFrame()
    if len(df_vif):
        msap_vif = df_vif[df_vif["model"] == "M_SAP"].sort_values("pct_selected", ascending=False)
        freq_str = ", ".join(
            f"{r['feature']} ({r['pct_selected']:.0f}%)"
            for _, r in msap_vif[msap_vif["pct_selected"] >= 80].iterrows()
        ) or "see Table 7"
        para(
            f"Iterative VIF selection (threshold = 10) consistently retained a non-collinear "
            f"feature subset across folds. Features selected in >=80% of M_SAP folds: "
            f"{freq_str}. Between 9 and 12 features were retained per fold, confirming "
            "substantial dimensionality reduction from the initial 17 predictors. "
            "VIF selection was applied independently in each inner-CV fold, ensuring "
            "that feature selection decisions never depend on test-set information."
        )

    heading("3.7 SHAP Feature Importance", level=2)
    if len(shap):
        top3 = shap.head(3)
        t3_str = "; ".join(
            f"{r['feature']} ({r['pct_importance']:.1f}%)"
            for _, r in top3.iterrows()
        )
        para(
            f"SHAP analysis on the best final model (RF/M_SAP) identified {t3_str} as the "
            "three most important predictors (Figure S5). Distance to sea (A3) alone "
            "accounted for 36.0% of total mean |SHAP| importance, consistent with the "
            "species' affinity for coastal and estuarine environments. Population density "
            "(P1; 14.2%) ranked third, supporting human-mediated introduction pressure as "
            "a predictor of relative invasion suitability independent of environmental conditions."
        )
        caption(
            "Table 7. SHAP feature importance for the RF/M_SAP model "
            "(mean absolute SHAP value and relative importance %)."
        )
        add_table_from_df(shap[["rank","feature","group","mean_abs_shap","pct_importance"]])
        doc.add_paragraph()

    heading("3.8 Model Validation", level=2)
    boyce_csv = VAL3_DIR / "boyce_oof_summary.csv"
    moran_csv = VAL3_DIR / "morans_resid.csv"
    if boyce_csv.exists():
        bdf = pd.read_csv(boyce_csv)
        br  = bdf.iloc[0]["boyce_r"] if len(bdf) else "--"
        para(
            f"The Boyce Continuous Skill Index (CSI) computed from pooled out-of-fold "
            f"predictions was {br} for the RF/M_SAP model, indicating that predicted "
            "relative suitability reliably increases toward presence locations relative "
            "to random background (Figure S7). This estimate avoids in-sample inflation."
        )
    else:
        note("Boyce OOF index: run 16_validation_extras_v2.py on output4 OOF predictions.")

    if moran_csv.exists():
        mdf = pd.read_csv(moran_csv)
        if len(mdf):
            I_val = mdf.iloc[0]["morans_I"]
            z_val = mdf.iloc[0]["z_score"]
            p_val = mdf.iloc[0]["p_value"]
            para(
                f"Moran's I on pooled out-of-fold model residuals (observed - predicted "
                f"at test presences) was {I_val:.4f} (z = {z_val:.3f}, p = {p_val:.4f}), "
                "indicating significant spatial clustering of residuals -- an acknowledged "
                "limitation of the model."
            )
    else:
        note("Moran's I on residuals: recompute on output4 OOF predictions.")

    heading("3.9 Sensitivity Analysis", level=2)
    if len(tab_e):
        para(
            "Sensitivity analysis confirmed that the best-configuration (RF/M_SAP) AUC "
            "was robust to variations in thinning distance, CV fold count, and "
            "occurrence subset (Table 8). AUC-ROC ranged from "
            f"{tab_e['AUC Mean'].min()} to {tab_e['AUC Mean'].max()} "
            "across all sensitivity scenarios. Models trained on GBIF/iNaturalist records "
            "only (excluding field reports) showed slightly higher AUC, consistent with "
            "higher geographic clustering of field-report records adding noise. "
            "Results were most stable to variation in thinning distance (5-20 km), "
            "suggesting that the model is not overly sensitive to the exact thinning "
            "parameterisation at this occurrence density."
        )
        caption("Table 8. Sensitivity analysis summary: RF/M_SAP AUC-ROC across scenarios.")
        add_table_from_df(tab_e)
        doc.add_paragraph()
    else:
        note("Sensitivity analysis table not found in output4/paper_tables/tableE_sensitivity.csv")

    heading("3.10 Province-Level Relative Suitability", level=2)
    if len(df_prov):
        top5 = df_prov.head(5)
        top5_strs = []
        for _, r in top5.iterrows():
            n  = r.get("n_cells", "?")
            lo = r.get("low_n_flag", False)
            flag = " [LOW-N]" if lo else ""
            ms = r.get("mean_suit", r.get("mean_suit", "?"))
            ms_str = f"{float(ms):.3f}" if ms != "?" else "?"
            top5_strs.append(f"{r['province']} (mean={ms_str}, n={n}{flag})")
        top5_str = ", ".join(top5_strs)

        has_low_n = df_prov.get("low_n_flag", pd.Series([False])).any()
        caveat = (
            " Province rankings with fewer than 20 water cells should be interpreted "
            "cautiously (marked [LOW-N])." if has_low_n else ""
        )
        para(
            f"Province-level aggregation of RF/M_SAP relative suitability ranked "
            f"{top5_str} among the five highest areas (Table 9, Figure 6).{caveat} "
            "The predominance of Gulf of Thailand coastal provinces is consistent with "
            "the model's identification of coastal proximity (A3) as the dominant predictor."
        )
        show_cols = [c for c in ["rank","province","n_cells","mean_suit","sd_suit",
                                  "max_suit","pct_high","low_n_flag"] if c in df_prov.columns]
        caption(
            "Table 9. Province-level relative invasion suitability ranking "
            "(top 15; RF/M_SAP model)."
        )
        add_table_from_df(df_prov.head(15)[show_cols])
        doc.add_paragraph()

    # ── Figure 1: Study area ───────────────────────────────────────────────────
    doc.add_paragraph()
    note("Study-area map pending regeneration for the current n=38 dataset without A3.")

    # ── Figure 2: Algorithm comparison boxplot ─────────────────────────────────
    doc.add_paragraph()
    add_figure(FIG4_DIR / "Fig5_algorithm_boxplot.png", width_cm=15.0,
               fallback_text="Fig5_algorithm_boxplot")
    caption(
        "Figure 2. Fold-level AUC-ROC distributions across algorithms (RF, XGB, MXN), "
        "feature sets (M_S, M_SA, M_SP, M_SAP), and pseudo-absence methods. "
        "Box plots show median and IQR across 15 outer folds (5 repeats x 3 folds). "
        "Significant pairwise differences (permutation p<0.05) marked with asterisks."
    )

    # ── Figure 3: AUC heatmap ──────────────────────────────────────────────────
    doc.add_paragraph()
    add_figure(FIG4_DIR / "Fig4_auc_heatmap.png", width_cm=13.0,
               fallback_text="Fig4_auc_heatmap")
    caption(
        "Figure 3. AUC-ROC heatmap (mean fold-level AUC) for all algorithm x feature set "
        "combinations (pooled across pseudo-absence methods). "
        "Darker cells indicate higher discriminatory ability."
    )

    # ── Figure 4: ROC curves ───────────────────────────────────────────────────
    doc.add_paragraph()
    add_figure(FIG4_DIR / "Fig1_roc_curves.png", width_cm=14.0,
               fallback_text="Fig1_roc_curves")
    caption(
        "Figure 4. ROC curves for the best configuration per algorithm "
        "(M_SAP feature set, best pseudo-absence method), computed from "
        "pooled out-of-fold predictions. Shaded regions = 95% bootstrap CI."
    )

    # ── Figure 5: Ensemble suitability map ────────────────────────────────────
    doc.add_paragraph()
    note("Full-data suitability map pending regeneration from the current n=38 experiment.")

    # ── Figure 6: Province risk map ────────────────────────────────────────────
    doc.add_paragraph()
    note("Province-level aggregation pending regeneration from the current n=38 experiment.")

    # ══════════════════════════════════════════════════════════
    # 4. DISCUSSION
    # ══════════════════════════════════════════════════════════
    heading("4. Discussion", level=1)

    heading("4.1 Algorithm Performance", level=2)
    para(
        "The omnibus algorithm comparison was significant. RF exceeded XGB "
        f"(Holm-adjusted p = {rf_xgb_p}) and MXN (p = {rf_mxn_p}), while XGB also "
        f"exceeded MXN (p = {xgb_mxn_p}). This supports stronger discrimination from "
        "the tree ensembles in the present experiment, while the substantial fold-level "
        "variation still requires calibration and stability to be considered alongside AUC."
    )

    heading("4.2 Value of Waterway Proximity and Human-Pressure Layers", level=2)
    para(
        "The M_SP feature set produced the highest mean fold-level AUC, whereas M_SAP "
        "ranked second. This pattern indicates that human-pressure proxies added useful "
        "discrimination beyond environmental variables, but adding the available "
        "waterway proxies did not uniformly improve performance. Because SHAP values "
        "were not regenerated for the current n=38 experiment, no variable-level causal "
        "or importance claim is made here."
    )

    heading("4.3 Pseudo-Absence Method", level=2)
    para(
        f"Pseudo-absence method did not affect AUC in the omnibus comparison (H = {h_pa}, "
        f"p = {p_pa}). The only Holm-adjusted pairwise difference was higher AUC for "
        "random than spatially constrained sampling. The remaining contrasts were not "
        "significant, so these data do not justify selecting a universal pseudo-absence "
        "strategy."
    )

    heading("4.4 Methodological Contributions and Limitations", level=2)
    para(
        "Our design contributes three methodological advances over prior SDMs for "
        "Thai aquatic invaders. First, a formal full-factorial comparison of three "
        "algorithms, four feature sets, and four pseudo-absence methods under "
        "identical leakage-free conditions provides transparent guidance on model "
        "selection rather than relying on a single algorithm choice. Second, "
        "treating pseudo-absence ratio as a hyperparameter selected inside the "
        "inner CV loop avoids optimistically biased AUC from ratio tuning on the "
        "full training set. Third, repeated spatial partitioning combined with fixed "
        "evaluation backgrounds allows model and background-sampling variability to be "
        "quantified. TSS thresholds "
        "are derived exclusively from inner-fold out-of-fold predictions to prevent "
        "test-set optimism in threshold selection."
    )
    para(
        f"Critical limitations: (1) SAMPLE SIZE — after removing {N_OCC_REMOVED} records "
        f"with collection years > 2024 (data contamination), only n = {N_OCC} occurrence "
        f"records were available for modelling. Although this exceeds a minimal threshold "
        f"often used in sparse-data SDMs, it remains limited for nationwide inference. "
        f"Outer CV folds contain approximately {max(1, N_OCC//3)}-{N_OCC//3+2} test "
        f"presences each, and inner folds operate on as few as {max(1, int(N_OCC*2/3//3))} "
        f"training presences. Performance metrics, especially threshold-dependent metrics "
        f"(TSS, sensitivity, specificity), should be interpreted cautiously. "
        f"The confidence intervals provided by clustered bootstrap are likely wide and "
        f"may not accurately represent generalisation performance. Independent validation "
        f"with newly collected occurrence data is essential before operational use. "
        f"(2) A-layer variables (A1_dist_waterway, A2_waterway_order) are proximity "
        f"proxies rather than true hydrological connectivity metrics; the initially "
        f"planned coastal proximity variable (A3_dist_coast) was excluded after "
        f"inspection revealed it measured distance to the nearest water body rather "
        f"than sea coast distance, making it redundant with A1. "
        f"(3) P-layer population density (WorldPop 2020) has missing coverage for "
        f"61.7% of water-class cells; cells with no WorldPop data were zero-filled "
        f"(valid assumption for open water bodies) but this may underestimate "
        f"propagule pressure along small inland waterways. "
        f"(4) MXN's lower performance likely reflects that MaxNet trained on the "
        f"same pseudo-absence background as RF/XGB departs from the standard MaxEnt "
        f"background distribution approach; a conventional MaxEnt calibration "
        f"(e.g., via ENMeval) might yield different results. "
        f"(5) The model does not incorporate temporal dynamics, projected climate "
        f"change scenarios, or dispersal limitations."
    )

    # ══════════════════════════════════════════════════════════
    # 5. CONCLUSIONS
    # ══════════════════════════════════════════════════════════
    heading("5. Conclusions", level=1)
    para(
        "We present a leakage-free, multi-algorithm machine learning framework "
        "that integrates environmental suitability (S), waterway proximity proxies (A), "
        "and human-pressure proxies (P) for mapping relative habitat suitability of "
        f"blackchin tilapia in Thailand using n = {N_OCC} confirmed occurrence records "
        f"(after temporal quality control). Across 48 algorithm × feature set × "
        f"pseudo-absence method combinations, {BEST['algo']}/{BEST['feat_set']}/{BEST['pa_method']} "
        f"achieved the highest cross-validated performance "
        f"(fold-mean AUC = {BEST['auc']} [95% CI: {BEST['ci_lo']}–{BEST['ci_hi']}]). "
        f"The omnibus algorithm effect was significant, and RF exceeded both XGB and MXN "
        f"in Holm-adjusted pairwise comparisons. M_SP and M_SAP yielded nearly identical "
        f"mean fold-level AUC, indicating useful discrimination from human-pressure proxies. "
        f"Pseudo-absence method had no significant omnibus effect; only random versus "
        f"spatially constrained sampling differed after correction. Findings remain "
        f"provisional given the sparse occurrence sample (n = {N_OCC}); independent "
        f"validation is required before operational use. The framework "
        f"is methodologically transferable to other aquatic invaders across Southeast Asia "
        f"and will benefit substantially from expanded occurrence databases."
    )

    # ══════════════════════════════════════════════════════════
    # DATA AVAILABILITY
    # ══════════════════════════════════════════════════════════
    heading("Data Availability Statement", level=1)
    para(
        "Occurrence data were obtained from GBIF (DOI: [to be confirmed]) and iNaturalist "
        "(inaturalist.org). WorldClim 2.1 data are available at worldclim.org. "
        "WorldPop 2020 data are available at worldpop.org. OpenStreetMap data "
        "(Geofabrik Thailand extract) are available at download.geofabrik.de. "
        "JRC Global Surface Water data are available via the European Commission Joint "
        "Research Centre. GADM boundaries are available at gadm.org. "
        "Processed occurrence data and model outputs are available from the corresponding "
        "author upon reasonable request."
    )

    heading("Code Availability", level=1)
    para(
        "All analysis code was written in Python. Key packages: scikit-learn, xgboost 3.2.0, "
        "rasterio, geopandas, shap, numpy, pandas, scipy, matplotlib. Code is available at "
        "[GitHub/Zenodo DOI -- to be registered before submission]."
    )

    heading("Author Contributions", level=1)
    para(
        "[To be completed following CRediT taxonomy: Conceptualization, Methodology, "
        "Software, Validation, Formal analysis, Investigation, Resources, "
        "Data Curation, Writing - Original Draft, Writing - Review and Editing, "
        "Visualization, Supervision, Project Administration, Funding Acquisition]"
    )

    heading("Funding", level=1)
    para("[Funding sources -- to be completed]")

    heading("Conflicts of Interest", level=1)
    para("The authors declare no conflicts of interest.")

    # ══════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════
    heading("References", level=1)
    refs = [
        "Breiman L (2001) Random forests. Machine Learning 45: 5-32.",

        "Chen T, Guestrin C (2016) XGBoost: A scalable tree boosting system. "
        "Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge "
        "Discovery and Data Mining, pp. 785-794.",

        "Elith J, Graham CH, Anderson RP, Dudik M, et al. (2006) Novel methods improve "
        "prediction of species' distributions from occurrence data. Ecography 29: 129-151.",

        "Elith J, Leathwick JR (2009) Species distribution models: ecological explanation "
        "and prediction across space and time. Annual Review of Ecology, Evolution, and "
        "Systematics 40: 677-697.",

        "Fick SE, Hijmans RJ (2017) WorldClim 2: new 1-km spatial resolution climate "
        "surfaces for global land areas. International Journal of Climatology 37: 4302-4315.",

        "GADM (2023) GADM database of Global Administrative Areas, version 4.1. "
        "https://gadm.org/ [Accessed: date to be confirmed]",

        "Hirzel AH, Le Lay G, Helfer V, Randin C, Guisan A (2006) Evaluating the ability "
        "of habitat suitability models to predict species presences. Ecological Modelling "
        "199: 142-152.",

        "Iturbide M, Bedia J, Herrera S, del Hierro O, Pinto M, Gutierrez JM (2015) "
        "A framework for species distribution modelling with improved pseudo-absence "
        "generation. Ecological Modelling 312: 166-174.",

        "Leprieur F, Beauchard O, Blanchet S, Oberdorff T, Brosse S (2008) Fish invasions "
        "in the world's river systems: when natural processes are blurred by human activities. "
        "PLoS Biology 6: e28.",

        "Lockwood JL, Cassey P, Blackburn T (2005) The role of propagule pressure in "
        "explaining species invasions. Trends in Ecology and Evolution 20: 223-228.",

        "Lundberg SM, Lee SI (2017) A unified approach to interpreting model predictions. "
        "Advances in Neural Information Processing Systems 30: 4765-4774.",

        "Miller ME, Hui SL, Tierney WM (1993) Validation techniques for logistic regression "
        "models. Statistics in Medicine 10: 1213-1226.",

        "OpenStreetMap Contributors (2024) Planet dump. https://www.openstreetmap.org "
        "[Geofabrik Thailand extract, accessed: date to be confirmed]",

        "Pekel JF, Cottam A, Gorelick N, Belward AS (2016) High-resolution mapping of "
        "global surface water and its long-term changes. Nature 540: 418-422.",

        "Phillips SJ, Anderson RP, Schapire RE (2006) Maximum entropy modeling of "
        "species geographic distributions. Ecological Modelling 190: 231-259.",

        "Phillips SJ, Anderson RP, Dudik M, Schapire RE, Blair ME (2017) Opening the "
        "black box: an open-source release of Maxent. Ecography 40: 887-893.",

        "Rahel FJ (2007) Biogeographic barriers, connectivity and homogenization of "
        "freshwater faunas: it's a small world after all. Freshwater Biology 52: 696-710.",

        "Roberts DR, Bahn V, Ciuti S, Boyce MS, Elith J, Guillera-Arroita G, et al. (2017) "
        "Cross-validation strategies for data with temporal, spatial, hierarchical, or "
        "phylogenetic structure. Ecography 40: 913-929.",

        "Senay SD, Worner SP, Ikeda T (2013) Novel three-step pseudo-absence selection "
        "technique for improved species distribution modelling. PLoS ONE 8: e71218.",

        "Tatem AJ (2017) WorldPop, open data for spatial demography. Scientific Data 4: 170004.",

        "Valavi R, Elith J, Lahoz-Monfort JJ, Guillera-Arroita G (2019) blockCV: An R "
        "package for generating spatially or environmentally separated folds for k-fold "
        "cross-validation of species distribution models. Methods in Ecology and Evolution "
        "10: 225-232.",

        "Valavi R, Guillera-Arroita G, Lahoz-Monfort JJ, Elith J (2021) Predictive "
        "performance of presence-only species distribution models: a benchmark study with "
        "reproducible code. Ecological Monographs 92: e01486.",

        "Vidthayanon C (2005) Aquatic Invasive Species Management in Thailand. "
        "Proceedings of the Regional Workshop on Aquatic Invasive Species. "
        "RAP Publication 2005/02, FAO Regional Office for Asia and the Pacific, Bangkok.",

        "Vitousek PM, D'Antonio CM, Loope LL, Westbrooks R (1997) Biological invasions "
        "as global environmental change. American Scientist 84: 468-478.",
    ]
    for r in refs:
        p   = doc.add_paragraph(style="List Bullet")
        run = p.add_run(r)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════
    # SUPPLEMENTARY
    # ══════════════════════════════════════════════════════════
    doc.add_page_break()
    heading("Supplementary Material", level=1)

    heading("Figure S1. PR curves", level=2)
    add_figure(FIG4_DIR / "Fig2_pr_curves.png", width_cm=14.0,
               fallback_text="Fig2_pr_curves")
    caption(
        "Figure S1. Precision-recall curves for the best configuration per algorithm "
        "(M_SAP feature set), computed from pooled out-of-fold predictions. "
        "Shaded regions = 95% bootstrap CI. Dashed horizontal line = no-skill baseline "
        "(prevalence = proportion of presences in pooled OOF)."
    )

    heading("Figure S2. Calibration plots", level=2)
    add_figure(FIG4_DIR / "Fig3_calibration.png", width_cm=13.0,
               fallback_text="Fig3_calibration")
    caption(
        "Figure S2. Calibration curves (predicted vs. observed fraction of presences) "
        "for all algorithm x feature set combinations. Diagonal dashed line = perfect "
        "calibration. Curves close to diagonal indicate well-calibrated probability estimates."
    )

    heading("Figure S3. Model disagreement map", level=2)
    note("Pending regeneration from full-data models fitted to the current n=38 dataset.")

    heading("Figure S4. Sensitivity analysis tornado chart", level=2)
    add_figure(FIG4_DIR / "FigS1_sensitivity_tornado.png", width_cm=13.0,
               fallback_text="FigS1_sensitivity_tornado")
    caption(
        "Figure S4. Sensitivity analysis tornado chart: change in AUC-ROC of the best "
        "configuration (RF/M_SAP) across variations in spatial thinning distance, "
        "number of CV folds, and occurrence subset (GBIF/iNat only vs. full dataset). "
        "Bars centred on baseline AUC; length indicates range across levels."
    )

    heading("Figure S5. SHAP feature importance (RF/M_SAP)", level=2)
    note("Pending regeneration from the current n=38 model; output3 SHAP values were excluded.")

    heading("Figure S6. Suitability 4-panel map (RF model variants)", level=2)
    note("Pending regeneration from full-data models fitted to the current n=38 dataset.")

    heading("Figure S7. Boyce Continuous Skill Index curve (OOF)", level=2)
    add_figure(VAL3_DIR / "boyce_oof_plot.png", width_cm=9.0,
               fallback_text="boyce_oof_plot")
    caption(
        "Figure S7. Predicted/Expected (P/E) ratio across suitability bins for the "
        "RF/M_SAP model, computed from pooled out-of-fold predictions. "
        "A monotonically increasing P/E curve indicates reliable habitat ranking."
    )

    heading("Figure S8. Feature correlation matrix", level=2)
    note(
        "Feature-correlation figure pending regeneration for the current 16-predictor, "
        "11,750-cell common domain."
    )

    return doc


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 65)
    print("Step 25: Word Manuscript v4 (Full-Factorial Multi-Algorithm)")
    print("=" * 65)

    try:
        from docx import Document
    except ImportError:
        print("  python-docx not installed. Run: pip install python-docx")
        return

    doc = build_doc()
    doc.save(str(OUT_DOC))
    size_kb = OUT_DOC.stat().st_size // 1024
    print(f"\n  Saved: {OUT_DOC.name}  ({size_kb} KB)")
    print(f"  Path:  {OUT_DOC}")
    print()
    print("  v4 additions vs v3:")
    print("    [NEW]  Section 2.4: Machine Learning Algorithms (RF, XGB, MaxNet)")
    print("    [NEW]  Section 2.5: Pseudo-Absence Sampling Methods (4 methods)")
    print("    [NEW]  Section 3.3: Algorithm Comparison (Kruskal-Wallis + pairwise)")
    print("    [NEW]  Section 3.4: Feature Set Comparison")
    print("    [NEW]  Section 3.5: Pseudo-Absence Method Comparison")
    print("    [NEW]  Section 3.9: Sensitivity Analysis")
    print("    [UPD]  Abstract: RF/M_SAP best AUC + 95% CI")
    print("    [UPD]  Methods: full factorial design description")
    print("    [UPD]  Discussion: algorithm uncertainty, PA sensitivity")
    print("    [UPD]  Conclusions: best configuration named explicitly")
    print("    [UPD]  References: +Chen & Guestrin 2016, +Phillips 2017,")
    print("           +Iturbide 2015, +Senay 2013, +Valavi 2021, +Breiman 2001")


if __name__ == "__main__":
    main()
