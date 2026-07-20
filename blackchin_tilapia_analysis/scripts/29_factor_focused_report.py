"""Build a factor-focused Paper 1 manuscript draft from current validated outputs."""

from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "output"
XAI = OUT / "factor_explainability"
FIG = XAI / "figures"
EVIDENCE = ROOT / "variable_evidence.csv"
REPORT = ROOT / "Paper1_Factor_Focused_Manuscript.docx"


def fmt(value, digits=3):
    return f"{float(value):.{digits}f}"


def add_table(doc, frame, max_rows=None):
    shown = frame.head(max_rows) if max_rows else frame
    table = doc.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    for i, column in enumerate(shown.columns):
        table.rows[0].cells[i].text = str(column)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if pd.isna(value) else str(value)
    return table


def add_picture(doc, name, width=15.5):
    path = FIG / name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width))


def main():
    summary = pd.read_csv(XAI / "feature_importance_summary.csv")
    groups = pd.read_csv(XAI / "factor_group_summary.csv")
    pooled = pd.read_csv(OUT / "pooled_oof_metrics.csv")
    folds = pd.read_csv(OUT / "experiment_fold_metrics.csv")
    evidence = pd.read_csv(EVIDENCE)
    best = pooled.sort_values("auc_roc", ascending=False).iloc[0]
    top = summary.iloc[0]
    top5 = summary.head(5)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Environmental, Waterway-Accessibility, and Human-Pressure Correlates "
        "of Relative Habitat Suitability for Invasive Blackchin Tilapia "
        "(Sarotherodon melanotheron) in Thailand"
    )
    run.bold = True
    run.font.size = Pt(15)
    doc.add_paragraph("[Authors and affiliations to be completed]").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph(f"Generated {date.today().isoformat()}").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Understanding which environmental and anthropogenic factors are associated "
        "with habitat suitability is central to surveillance of invasive aquatic fishes. "
        "We evaluated environmental suitability (S), waterway-accessibility proxies (A), "
        "and human-pressure proxies (P) for blackchin tilapia in Thailand using 38 "
        "quality-controlled occurrences. A leakage-free nested spatial cross-validation "
        "experiment compared 48 algorithm, predictor-set, and pseudo-absence combinations. "
        f"The best configuration was {best['algo']}/{best['feat_set']}/"
        f"{best['pa_method']} (pooled out-of-fold AUC={fmt(best['auc_roc'])}, "
        f"95% clustered-bootstrap CI {fmt(best['auc_roc_ci_lo'])}-"
        f"{fmt(best['auc_roc_ci_hi'])}). Held-out permutation importance, SHAP, and "
        "accumulated local effects were then used to interpret fitted associations. "
        f"The strongest held-out factor was {top['label']} "
        f"(mean AUC decrease={fmt(top['permutation_auc_decrease_mean'])}), while "
        "temperature-related predictors, elevation, population density, and waterway "
        "order collectively dominated model interpretation. These are predictive "
        "associations rather than causal effects. The resulting aquatic-network map "
        "identifies relative suitability and should not be interpreted as calibrated "
        "invasion probability."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Blackchin tilapia is a euryhaline West African cichlid associated with lagoons "
        "and estuaries (Pauly 1976). Its spread in Thailand creates a need to identify "
        "environmental conditions and human-associated pathways linked to suitable "
        "habitat. Most species-distribution studies emphasize predictive accuracy, but "
        "management additionally requires transparent evidence about which factors are "
        "consistently informative. We therefore frame this study around correlates of "
        "relative habitat suitability rather than causal drivers or invasion probability."
    )
    doc.add_paragraph(
        "The study asks: (1) Do waterway-accessibility and human-pressure proxies add "
        "information beyond environmental variables? (2) Which individual factors are "
        "stable across spatially independent test folds? (3) What fitted response shapes "
        "are suggested by the best validated model?"
    )

    doc.add_heading("2. Materials and Methods", level=1)
    doc.add_heading("2.1 Occurrences and predictors", level=2)
    doc.add_paragraph(
        "The occurrence dataset contained 38 records dated no later than 2024. Predictor "
        "families comprised ten bioclimatic/topographic variables (S), distance to "
        "waterway and waterway order (A), and population density, road distance, road "
        "density, and urban distance (P). A3 distance-to-coast was excluded because its "
        "distribution indicated distance to the nearest water body rather than true sea "
        "coast distance. P variables are described as human-pressure proxies, not direct "
        "measurements of propagule pressure."
    )

    doc.add_heading("2.2 Validation and model selection", level=2)
    doc.add_paragraph(
        "Three algorithms, four predictor sets, and four pseudo-absence strategies were "
        "compared using five repetitions of balanced three-fold spatial cross-validation "
        "(12-13 test presences per fold). Feature selection, pseudo-absence ratio, "
        "hyperparameter tuning, and the TSS threshold were estimated within training "
        "data. Technical details and complete factorial results are retained in the "
        "supplementary output; the main text focuses on ecological interpretation."
    )

    doc.add_heading("2.3 Factor interpretation", level=2)
    doc.add_paragraph(
        "For the best configuration, all 15 outer-fold models were reconstructed exactly; "
        "their AUC values matched the reported experiment values. Permutation importance "
        "was calculated only on each held-out outer test fold as the decrease in AUC after "
        "30 independent permutations. Fold-level bootstrap confidence intervals used "
        "2,000 resamples. SHAP values summarized model contributions, and accumulated "
        "local effects (ALE) described fitted response shapes under correlated predictors. "
        "A variable absent from a fold-specific VIF-selected model received zero ensemble "
        "importance for that fold. None of these analyses establishes causality."
    )

    doc.add_heading("2.4 Aquatic-network mapping", level=2)
    doc.add_paragraph(
        "The full-country prediction was masked using OpenStreetMap water polygons and "
        "river, canal, and stream features. Linear waterways were buffered by 500 m and "
        "rasterized with all-touched inclusion to match the approximately 4.6-km predictor "
        "grid. The continuous map is presented as relative habitat suitability. Binary "
        "classification is not emphasized because no independent validation dataset was "
        "available."
    )

    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Predictive performance", level=2)
    doc.add_paragraph(
        f"The selected {best['algo']}/{best['feat_set']}/{best['pa_method']} configuration "
        f"had pooled out-of-fold AUC={fmt(best['auc_roc'])}, PR-AUC="
        f"{fmt(best['pr_auc'])}, TSS={fmt(best['tss'])}, and Brier score="
        f"{fmt(best['brier'])}. Across all folds, RF mean AUC was "
        f"{fmt(folds.loc[folds.algo == 'RF', 'auc_roc'].mean())}, compared with "
        f"{fmt(folds.loc[folds.algo == 'XGB', 'auc_roc'].mean())} for XGB and "
        f"{fmt(folds.loc[folds.algo == 'MXN', 'auc_roc'].mean())} for MXN."
    )

    doc.add_heading("3.2 Individual factors", level=2)
    names = ", ".join(
        f"{row.label} ({fmt(row.permutation_auc_decrease_mean)})"
        for row in top5.itertuples())
    doc.add_paragraph(
        "The five largest held-out AUC decreases were observed for " + names + ". "
        "Population density was selected in every fold, as was waterway order. "
        "Positive or negative SHAP-direction correlations describe the average fitted "
        "direction only and should be interpreted together with the non-linear ALE curves."
    )
    importance_table = summary.head(12)[[
        "label", "group", "selection_frequency",
        "permutation_auc_decrease_mean", "permutation_ci_lo",
        "permutation_ci_hi", "shap_direction_rho"]].copy()
    importance_table.columns = [
        "Factor", "Group", "Selection frequency", "AUC decrease",
        "95% CI low", "95% CI high", "SHAP direction rho"]
    for col in importance_table.columns[2:]:
        importance_table[col] = importance_table[col].map(lambda x: fmt(x))
    add_table(doc, importance_table)
    add_picture(doc, "Fig_factor_permutation_importance.png")

    doc.add_heading("3.3 Predictor-family contribution", level=2)
    s_value = groups.loc[
        groups.group.str.startswith("S:"), "permutation_auc_decrease"].iloc[0]
    p_value = groups.loc[
        groups.group.str.startswith("P:"), "permutation_auc_decrease"].iloc[0]
    a_value = groups.loc[
        groups.group.str.startswith("A:"), "permutation_auc_decrease"].iloc[0]
    doc.add_paragraph(
        f"Summed permutation importance was greatest for S ({fmt(s_value)}), followed "
        f"by P ({fmt(p_value)}) and A ({fmt(a_value)}). This supports a dominant "
        "environmental signal with additional information from human-pressure and "
        "waterway-accessibility proxies."
    )
    add_picture(doc, "Fig_SHAP_top_factors.png")

    doc.add_heading("3.4 Fitted response shapes", level=2)
    doc.add_paragraph(
        "ALE curves indicated non-linear responses and substantial model-to-model "
        "uncertainty for several factors. Elevation generally showed a negative fitted "
        "association, whereas population density and waterway order showed positive "
        "average SHAP-direction correlations. Temperature responses were non-monotonic, "
        "which is consistent with a bounded environmental niche rather than a simple "
        "linear optimum."
    )
    add_picture(doc, "Fig_ALE_top_factors.png")

    doc.add_heading("3.5 Spatial pattern", level=2)
    doc.add_paragraph(
        "The OSM-masked raster retained predictions along mapped water bodies and the "
        "river-canal-stream network. The map is suitable for relative prioritization, "
        "not for estimating occurrence probability or proving current presence."
    )
    add_picture(doc, "Fig_aquatic_suitability_map.png", width=12.5)

    doc.add_heading("4. Discussion", level=1)
    doc.add_paragraph(
        "Temperature variability and warm-season temperature were among the most stable "
        "predictors, consistent with physiological constraints on ectothermic fishes. "
        "Elevation likely captures the lowland and coastal distribution of connected "
        "aquatic habitats. Population density added predictive information but may "
        "combine introduction pressure, habitat modification, observation intensity, "
        "and reporting bias. It must therefore not be interpreted as a direct causal "
        "measure of fish release."
    )
    doc.add_paragraph(
        "Waterway order was selected in every fold, supporting the relevance of network "
        "position and aquatic accessibility. Its smaller permutation effect than several "
        "environmental factors suggests that accessibility modifies, rather than replaces, "
        "environmental suitability. Distance to waterway was less stable, possibly because "
        "the analysis domain was already restricted toward aquatic environments."
    )
    doc.add_paragraph(
        "Important omissions remain. Salinity, water temperature, pH, dissolved oxygen, "
        "and turbidity were not available as spatially complete national layers at the "
        "required resolution. Climatic variables therefore act as indirect proxies for "
        "some aquatic processes. Future field sampling should prioritize these water-quality "
        "measurements and independent presence-absence validation."
    )

    doc.add_heading("5. Limitations", level=1)
    doc.add_paragraph(
        "The dataset remains small (n=38), no independent validation dataset was available, "
        "and pseudo-absence evaluation does not substitute for confirmed absences. The "
        "waterway and human-pressure layers are proxies. SHAP and ALE explain fitted model "
        "behavior, not biological causation. The 4.6-km grid is coarse relative to narrow "
        "canals, and OSM completeness is spatially uneven. These constraints preclude "
        "labelling the map as calibrated invasion risk."
    )

    doc.add_heading("6. Conclusions", level=1)
    doc.add_paragraph(
        "Blackchin tilapia relative habitat suitability in Thailand was most strongly "
        "associated with environmental factors, particularly temperature variability, "
        "warm-season temperature, and low-elevation settings. Population density and "
        "waterway order supplied additional predictive information, supporting the use "
        "of human-pressure and accessibility proxies alongside environmental layers. "
        "The study provides a leakage-aware explanatory framework, while independent "
        "surveys and direct water-quality layers remain necessary for operational decisions."
    )

    doc.add_heading("Table S1. Evidence supporting predictor selection", level=1)
    evidence_show = evidence[[
        "variable_group", "variables", "ecological_rationale",
        "evidence_scope", "source_title", "journal_or_source",
        "doi_or_url", "manuscript_status"]].copy()
    evidence_show.columns = [
        "Group", "Variables", "Rationale", "Evidence scope", "Source",
        "Journal/source", "DOI/URL", "Use"]
    add_table(doc, evidence_show)

    doc.add_heading("Data Availability", level=1)
    doc.add_paragraph(
        "The analysis folder contains fold metrics, out-of-fold predictions, factor "
        "importance tables, SHAP values, ALE curves, figures, raster metadata, and "
        "reproducible scripts. GBIF download DOI and author/affiliation details must be "
        "completed before submission."
    )

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
